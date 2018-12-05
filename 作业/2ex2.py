

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# 设置0级标题并居中
document.add_heading('个人简历', level=0).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('基本信息', level=1)
# 建立1x3表格并放置内容
table = document.add_table(rows=1, cols=3)
cell1 = table.cell(0, 0)
cell1.text = """
姓名：项义军
民族：汉
电话：18896656769
邮箱：wind98@gmail.com
"""
cell2 = table.cell(0,1)
cell2.text = """
出生年月：1998.5
政治面貌：共青团员
学历：本科
毕业院校：江苏大学
"""










# 导入照片
cell3 = table.cell(0,2)
paragraph = cell3.paragraphs[0]
run = paragraph.add_run()
run.add_picture('1.jpg', width=Inches(1.25), height=Inches(1.4))

document.add_heading('教育背景', level=1)
# 建立1x3表格存放内容
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '2016.09-2020.06'
hdr_cells[1].text = '江苏大学'
hdr_cells[2].text = '复合材料'

paragraph = document.add_paragraph('核心课程：工程图学、电子电工学、机械设计基础、材料力学性能、材料物理性能\n')
paragraph.add_run('复合材料结构设计、复合材料工艺及设备、材料科学基础、复合材料原理')

document.add_heading('技能证书', level=1)
paragraph = document.add_paragraph('江苏省计算机二级考试证书、英语四六级证书')
document.add_heading('兴趣爱好', level=1)
paragraph = document.add_paragraph('轮滑、中国古代历史、慢跑....')

document.save('简历.docx')